import pymupdf
from googletrans import Translator
from docx import Document
import os
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


shrinking_factor_x = 0.11
shrinking_factor_y = 0.5
path = "/home/user/path_to_pdf.pdf" # change this line or uncomment the next one
#path = input("path_to_pdf: ")
file_name = os.path.basename(path)
file_name = os.path.splitext(file_name)[0]

doc = pymupdf.open(path)

annots_text = [] # contains highlighted text and the page
for page in doc:
    for annot in page.annots():
        if annot.type[1] == ('Highlight'):
           # print(f"Annotation on page: {page.number} with type: {annot.type} and rect: {annot.rect}")

            # shrinking the rect inwardly to prevent words on top to also be returned
            rect = annot.rect
            width = rect.x1 - rect.x0
            height = rect.y1 - rect.y0
            
            # if the highlighting goes through multiple lines
            if height > 30:
                text_blocks = page.get_text("dict")["blocks"]
                H_text = ""

                for block in text_blocks:
                    for line in block["lines"]:
                        for span in line["spans"]:
                            span_rect = span["bbox"]
                        
                            # Check if span is within the highlighted area
                            if (span_rect[0] >= rect.x0 and span_rect[2] <= rect.x1 and
                                span_rect[1] >= rect.y0 and span_rect[3] <= rect.y1):
                                H_text += span["text"] + " "
                text = H_text.strip(" ,.“”\"'")

            else:
                factor_x = width * shrinking_factor_x / 2
                factor_y = height * shrinking_factor_y / 2
                rect.x0 += factor_x
                rect.x1  -= factor_x
                rect.y0 += factor_y
                rect.y1 -= factor_y

                text = page.get_textbox(rect).strip(" ,.“”\"'")

            current_page = page.number
            #annots_text.append([text, current_page])
            annots_text.append([text])

doc.close
def translate_text(texts, src_lang="en", dest_lang="de"):
    translator = Translator()
    translated_text = ""
    for text in texts:
        translation = translator.translate(text, src=src_lang, dest=dest_lang)
        translated_text += (translation.text)
    return translated_text

for words in annots_text:
    word = [words[0]]
    translation = translate_text(word)
    middle_index = len(words) // 2 if len(words) > 1 else 1
    words.insert(middle_index, translation)

document = Document()


heading = document.add_heading(level=1)

# Create a hyperlink
part = document.part
path_with_page = f"{path}#page=100"
r_id = part.relate_to(path_with_page, 'hyperlink', is_external=True)
hyperlink = OxmlElement('w:hyperlink')
hyperlink.set(qn('r:id'), r_id)

# Create a run and add it to the hyperlink
new_run = OxmlElement('w:r')
rPr = OxmlElement('w:rPr')
new_run.append(rPr)
text = OxmlElement('w:t')
text.text = file_name
new_run.append(text)
hyperlink.append(new_run)

# Append the hyperlink to the heading
heading._element.append(hyperlink)

document.add_paragraph("")
table = document.add_table(rows=0, cols=2)

for english, german in annots_text:
    row_cells = table.add_row().cells
    row_cells[0].text = english
    row_cells[1].text = german

document.save(f"/home/user/vocabs/{file_name}.docx") #make sure to create a dir called vocabs
print(file_name)

